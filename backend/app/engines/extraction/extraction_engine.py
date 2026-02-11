"""
Multi-format document extraction engine.
Supports PDF, DOCX, XLSX, images with OCR fallback.
"""
import os
from typing import Dict, List, Optional, Tuple
from pathlib import Path
import pytesseract
from PIL import Image
import fitz  # PyMuPDF
import pdfplumber
from docx import Document as DocxDocument
from openpyxl import load_workbook
from loguru import logger

from app.core.config import settings


class ExtractionEngine:
    """Multi-format document extraction with OCR support."""
    
    def __init__(self):
        """Initialize extraction engine."""
        if settings.TESSERACT_PATH:
            pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_PATH
        
        logger.info("Extraction engine initialized")
    
    async def extract_document(self, file_path: str) -> Dict:
        """
        Extract content from document based on file type.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing extracted text, tables, metadata, and quality metrics
        """
        file_ext = Path(file_path).suffix.lower()
        
        try:
            if file_ext == '.pdf':
                return await self._extract_pdf(file_path)
            elif file_ext == '.docx':
                return await self._extract_docx(file_path)
            elif file_ext == '.xlsx':
                return await self._extract_xlsx(file_path)
            elif file_ext in ['.png', '.jpg', '.jpeg']:
                return await self._extract_image(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
        
        except Exception as e:
            logger.error(f"Extraction failed for {file_path}: {e}")
            raise
    
    async def _extract_pdf(self, file_path: str) -> Dict:
        """Extract content from PDF files."""
        logger.info(f"Extracting PDF: {file_path}")
        
        result = {
            "text": "",
            "tables": [],
            "metadata": {},
            "page_count": 0,
            "has_images": False,
            "ocr_used": False,
            "ocr_confidence": None
        }
        
        # Try digital extraction first with PyMuPDF
        try:
            doc = fitz.open(file_path)
            result["page_count"] = len(doc)
            result["metadata"] = doc.metadata
            
            text_parts = []
            for page_num, page in enumerate(doc):
                # Extract text
                page_text = page.get_text()
                text_parts.append(page_text)
                
                # Check for images
                if page.get_images():
                    result["has_images"] = True
            
            result["text"] = "\n\n".join(text_parts)
            doc.close()
            
            # If no text extracted, try OCR
            if not result["text"].strip():
                logger.info("No text found, attempting OCR")
                ocr_result = await self._ocr_pdf(file_path)
                result["text"] = ocr_result["text"]
                result["ocr_used"] = True
                result["ocr_confidence"] = ocr_result["confidence"]
        
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed, trying OCR: {e}")
            ocr_result = await self._ocr_pdf(file_path)
            result["text"] = ocr_result["text"]
            result["ocr_used"] = True
            result["ocr_confidence"] = ocr_result["confidence"]
        
        # Extract tables with pdfplumber
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    tables = page.extract_tables()
                    if tables:
                        result["tables"].extend(tables)
        except Exception as e:
            logger.warning(f"Table extraction failed: {e}")
        
        logger.info(f"PDF extraction complete: {len(result['text'])} chars, {len(result['tables'])} tables")
        return result
    
    async def _ocr_pdf(self, file_path: str) -> Dict:
        """Perform OCR on PDF pages."""
        from pdf2image import convert_from_path
        
        logger.info(f"Performing OCR on PDF: {file_path}")
        
        # Convert PDF to images
        images = convert_from_path(file_path, dpi=settings.OCR_DPI)
        
        text_parts = []
        confidences = []
        
        for i, image in enumerate(images):
            # Perform OCR
            ocr_data = pytesseract.image_to_data(
                image,
                lang=settings.OCR_LANGUAGE,
                output_type=pytesseract.Output.DICT
            )
            
            # Extract text
            page_text = pytesseract.image_to_string(image, lang=settings.OCR_LANGUAGE)
            text_parts.append(page_text)
            
            # Calculate confidence
            conf_values = [int(conf) for conf in ocr_data['conf'] if conf != '-1']
            if conf_values:
                confidences.append(sum(conf_values) / len(conf_values))
        
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0
        
        return {
            "text": "\n\n".join(text_parts),
            "confidence": avg_confidence
        }
    
    async def _extract_docx(self, file_path: str) -> Dict:
        """Extract content from DOCX files."""
        logger.info(f"Extracting DOCX: {file_path}")
        
        doc = DocxDocument(file_path)
        
        # Extract text
        text_parts = [para.text for para in doc.paragraphs]
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        # Extract metadata
        metadata = {
            "author": doc.core_properties.author,
            "created": str(doc.core_properties.created),
            "modified": str(doc.core_properties.modified),
            "title": doc.core_properties.title
        }
        
        return {
            "text": "\n".join(text_parts),
            "tables": tables,
            "metadata": metadata,
            "ocr_used": False
        }
    
    async def _extract_xlsx(self, file_path: str) -> Dict:
        """Extract content from Excel files."""
        logger.info(f"Extracting XLSX: {file_path}")
        
        wb = load_workbook(file_path, data_only=True)
        
        tables = []
        text_parts = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            
            # Extract as table
            sheet_data = []
            for row in sheet.iter_rows(values_only=True):
                sheet_data.append([str(cell) if cell is not None else "" for cell in row])
            
            tables.append({
                "sheet_name": sheet_name,
                "data": sheet_data
            })
            
            # Also create text representation
            text_parts.append(f"Sheet: {sheet_name}")
            for row in sheet_data:
                text_parts.append(" | ".join(row))
        
        return {
            "text": "\n".join(text_parts),
            "tables": tables,
            "metadata": {"sheet_count": len(wb.sheetnames)},
            "ocr_used": False
        }
    
    async def _extract_image(self, file_path: str) -> Dict:
        """Extract text from images using OCR."""
        logger.info(f"Extracting image with OCR: {file_path}")
        
        image = Image.open(file_path)
        
        # Perform OCR
        text = pytesseract.image_to_string(image, lang=settings.OCR_LANGUAGE)
        
        # Get confidence
        ocr_data = pytesseract.image_to_data(
            image,
            lang=settings.OCR_LANGUAGE,
            output_type=pytesseract.Output.DICT
        )
        
        conf_values = [int(conf) for conf in ocr_data['conf'] if conf != '-1']
        avg_confidence = sum(conf_values) / len(conf_values) if conf_values else 0
        
        return {
            "text": text,
            "tables": [],
            "metadata": {
                "width": image.width,
                "height": image.height,
                "format": image.format
            },
            "ocr_used": True,
            "ocr_confidence": avg_confidence
        }
